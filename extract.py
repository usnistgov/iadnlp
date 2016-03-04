#!/usr/bin/env python3.5

import argparse
import os
import os.path

import docx
import docx.enum.text
import makecloud
import statistics

# overrides specifies the respondent for each named interview

overrides = {"5819565_P12_EXP DC 107.docx": "S2",
             "5869145_P2_MID DC 201 possibly move to GP.docx": "S2",
             "MW303.docx": "S2",
             "MW308.docx": "S2",
             "P7 A.docx": None,
             "P7 Total NI 109.docx": None,
             }
cache = False

def process(fname):
    """Process file fname.docx and create a transcript file fname.txt that has just the text of
    the respondent. The respondent is determined by linguistic analysis---it's the person who said
    the most, unless there is an overrride."""

    # Remove trans filenames
    fbase = os.path.basename(fname)
    if fbase.startswith("~") or not fname.endswith(".docx"):
        return
    if "NOTES" in fname.upper():
        return
    transcript_fname = fname.replace(".docx", ".txt")

    # don't create a transcript that already exists...
    transcript_exists = os.path.exists(transcript_fname) and os.path.getsize(transcript_fname) > 0 
    if transcript_exists and (not args.force) and (not args.stats):
        return

    d = docx.Document(fname)
    # Make sure transcript file is properly encoded
    assert d.tables[0].rows[0].cells[0].text == 'Date:'
    date = d.tables[0].rows[0].cells[1].text
    print("{} ({}) ".format(os.path.basename(fname),date),end='   ')

    # Compute stats?
    if args.stats:
        times = []
        for table in d.tables[1:]:
            for row in table.rows:
                try:
                    c0 = row.cells[0].text
                    if c0[0]=='S' and c0[2]==' ' and c0[5]==':':
                        time = c0[3:8]
                        times.append(int(time[0:2])*60+int(time[3:5]))
                except IndexError:
                    continue
        return max(times)
            

    if fbase in overrides:
        print("{} is an overrride".format(fbase))
        os.unlink(transcript_fname)

    if cache and transcript_exists:
        return

    # Remember output for each speaker.
    speakers = {}
    errors = []
    for table in d.tables[1:]:
        for row in table.rows:
            try:
                c0 = row.cells[0].text
                c1 = row.cells[1].text
            except IndexError:
                continue
            speaker = c0[0:2]
            when    = c0[4:]
            if speaker:
                if speaker not in speakers: speakers[speaker] = []
                speakers[speaker].append(c1)
                continue
            if c1=='[silence]':
                continue
            errors.append("** unknown c0: {}  c1: {}".format(c0,c1))

    # After examining several approaches for distinguishing the respondent
    # (vocab, # of question marks, etc),
    # we decided to go with the longest text as the respondent.

    if fbase in overrides:
        respondent = overrides[fbase]
        rtext = "\n".join(speakers[respondent])
        print("{} override. Speaker is {}".format(fbase, respondent))
    else:
        text_per_speaker = ["\n".join(speakers[speaker]) for speaker in speakers.keys()]
        (length, rtext) = max([(len(text), text) for text in text_per_speaker])

    if not transcript_exists and not args.force:
        with open(transcript_fname, "w") as f:
            f.write(rtext)


def process_root(fname):
    if os.path.isfile(fname):
        process(fname)
    if os.path.isdir(fname):
        for (dirpath, dirnames, filenames) in os.walk(fname):
            for filename in filenames:
                if filename.endswith(".docx"):
                    process(os.path.join(dirpath, filename))

if __name__=="__main__":
    try:
        os.mkdir("transcripts")
    except FileExistsError:
        pass
    parser = argparse.ArgumentParser("Analyze DOCX transcript files from TranscribeMe!")
    parser.add_argument("files",help="Files or directories to analyze",nargs="*")
    parser.add_argument("--force",help="force output",action='store_true')
    parser.add_argument("--stats",help="Calculate stats",action='store_true')
    args = parser.parse_args()
    rets = []
    
    if args.stats:
        for (label, files) in makecloud.SOURCES.items():
            times = []
            print("======================")
            print("{}".format(label))
            print("")
            for fname in files:
                seconds = process(fname)
                print(seconds)
                if seconds:
                    times.append(seconds)
            print("   mean: {}  median: {}  stddev: {}".
                  format(statistics.mean(times)//1,statistics.median(times)//1,statistics.stdev(times)//1))
            mean = int(statistics.mean(times))
            print("   mean: {}:{:02}".format(mean//60,mean%60))
        exit(0)

            
        

    for fname in args.files if args.files else makecloud.ALL_DOCX:
        ret = process(fname)
        rets.append(ret)

